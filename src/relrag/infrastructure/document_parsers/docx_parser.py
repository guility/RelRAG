"""Parser for .docx (Office Open XML Word)."""

import io
from pathlib import Path

from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from relrag.domain.value_objects import PropertyType
from relrag.infrastructure.document_parsers.base import ParseResult
from relrag.infrastructure.document_parsers.metadata_keys import (
    PARSER_KEY_TO_CANONICAL,
    normalize_value_for_storage,
)


def _map_metadata(core_props: object) -> dict[str, tuple[str, PropertyType]]:
    """Map docx core properties to canonical keys."""
    result: dict[str, tuple[str, PropertyType]] = {}
    cp = core_props
    raw: dict[str, str | object | None] = {}
    for name in ("title", "subject", "author", "created", "modified", "last_modified_by", "language"):
        if hasattr(cp, name):
            raw[name] = getattr(cp, name)
    for parser_key, val in raw.items():
        if val is None or val == "":
            continue
        canon = PARSER_KEY_TO_CANONICAL.get(parser_key)
        if not canon:
            continue
        canonical_key, ptype = canon
        if ptype == PropertyType.DATE and hasattr(val, "isoformat"):
            str_val = val.isoformat()
        else:
            str_val = normalize_value_for_storage(val)
        result[canonical_key] = (str_val, ptype)
    return result


def parse_docx(data: bytes, filename: str | None = None) -> ParseResult:
    """Extract text and metadata from .docx bytes."""
    try:
        doc = DocxDocument(io.BytesIO(data))
    except (PackageNotFoundError, BadZipFile) as e:
        raise ValueError("Invalid or corrupted docx file") from e
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                tables_text.append(" ".join(cells))
    text = "\n\n".join(paragraphs)
    if tables_text:
        text += "\n\n" + "\n".join(tables_text)
    properties = _map_metadata(doc.core_properties)
    if filename:
        p = Path(filename)
        properties["source_file_name"] = (normalize_value_for_storage(p.name), PropertyType.STRING)
        properties["source_file_type"] = (normalize_value_for_storage("docx"), PropertyType.STRING)
    return ParseResult(text=text or " ", properties=properties)
